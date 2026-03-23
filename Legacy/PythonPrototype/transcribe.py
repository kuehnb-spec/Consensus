#!/usr/bin/env python3
"""
Hearing Transcriber - Local transcription with speaker diarization.
Uses WhisperX (OpenAI Whisper + PyAnnote) for fully local processing.
"""

import argparse
import json
import os
import sys
import time
from datetime import timedelta
from pathlib import Path

import whisperx
import torch


def format_timestamp(seconds: float) -> str:
    """Convert seconds to HH:MM:SS format."""
    td = timedelta(seconds=seconds)
    total_seconds = int(td.total_seconds())
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def transcribe_audio(
    audio_path: str,
    model_size: str = "large-v3",
    device: str = None,
    compute_type: str = None,
    hf_token: str = None,
    min_speakers: int = None,
    max_speakers: int = None,
    language: str = "en",
    batch_size: int = 16,
) -> dict:
    """
    Transcribe an audio file with speaker diarization.

    Args:
        audio_path: Path to the audio file.
        model_size: Whisper model size (tiny, base, small, medium, large-v3).
        device: Device to use (cuda, mps, cpu). Auto-detected if None.
        compute_type: Compute type (float16, int8, float32). Auto-selected if None.
        hf_token: HuggingFace token for PyAnnote models.
        min_speakers: Minimum expected speakers (helps accuracy).
        max_speakers: Maximum expected speakers (helps accuracy).
        language: Language code (default: en).
        batch_size: Batch size for transcription.

    Returns:
        dict with keys: segments, speakers, audio_path, duration
    """

    # Auto-detect device
    # Note: WhisperX uses CTranslate2 which does NOT support MPS.
    # On Apple Silicon we must use CPU for transcription/alignment.
    # PyAnnote (diarization) can use MPS, so we track both.
    if device is None:
        if torch.cuda.is_available():
            device = "cuda"
        else:
            device = "cpu"

    # Diarization device — PyAnnote supports MPS
    diarize_device = device
    if device == "cpu" and hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
        diarize_device = "mps"

    # Auto-select compute type
    if compute_type is None:
        if device == "cuda":
            compute_type = "float16"
        else:
            compute_type = "int8"

    print(f"Using device: {device} | diarization: {diarize_device} | compute: {compute_type} | model: {model_size}")

    # Resolve HF token
    if hf_token is None:
        hf_token = os.environ.get("HF_TOKEN")
    if not hf_token:
        print(
            "WARNING: No HuggingFace token provided. Speaker diarization requires a token."
        )
        print("  Set HF_TOKEN env var or pass --hf-token argument.")
        print("  Get your token at: https://huggingface.co/settings/tokens")
        print(
            "  You also need to accept the model agreements at:"
        )
        print(
            "    https://huggingface.co/pyannote/segmentation-3.0"
        )
        print(
            "    https://huggingface.co/pyannote/speaker-diarization-3.1"
        )

    # ── Step 1: Load audio ──
    print(f"\n[1/4] Loading audio: {audio_path}")
    audio = whisperx.load_audio(audio_path)
    duration = len(audio) / 16000  # whisperx loads at 16kHz
    print(f"       Duration: {format_timestamp(duration)}")

    # ── Step 2: Transcribe with Whisper ──
    print(f"\n[2/4] Transcribing with Whisper ({model_size})...")
    t0 = time.time()
    model = whisperx.load_model(
        model_size,
        device,
        compute_type=compute_type,
        language=language,
    )
    result = model.transcribe(audio, batch_size=batch_size, language=language)
    elapsed = time.time() - t0
    print(f"       Transcription complete in {elapsed:.1f}s")

    # Free GPU memory
    del model
    if device == "cuda":
        torch.cuda.empty_cache()

    # ── Step 3: Align timestamps ──
    print(f"\n[3/4] Aligning word-level timestamps...")
    t0 = time.time()
    model_a, metadata = whisperx.load_align_model(
        language_code=language, device=device
    )
    result = whisperx.align(
        result["segments"], model_a, metadata, audio, device, return_char_alignments=False
    )
    elapsed = time.time() - t0
    print(f"       Alignment complete in {elapsed:.1f}s")

    del model_a
    if device == "cuda":
        torch.cuda.empty_cache()

    # ── Step 4: Speaker diarization ──
    if hf_token:
        print("\n[4/4] Identifying speakers (diarization)...")
        t0 = time.time()

        diarize_kwargs = {}
        if min_speakers is not None:
            diarize_kwargs["min_speakers"] = min_speakers
        if max_speakers is not None:
            diarize_kwargs["max_speakers"] = max_speakers

        diarize_model = whisperx.DiarizationPipeline(
            use_auth_token=hf_token, device=diarize_device
        )
        diarize_segments = diarize_model(audio, **diarize_kwargs)
        result = whisperx.assign_word_speakers(diarize_segments, result)

        elapsed = time.time() - t0
        num_speakers = len(set(
            seg.get("speaker", "UNKNOWN") for seg in result["segments"]
        ))
        print(f"       Diarization complete in {elapsed:.1f}s")
        print(f"       Detected {num_speakers} speakers")
    else:
        print("\n[4/4] Skipping diarization (no HF token)")

    return {
        "segments": result["segments"],
        "audio_path": str(audio_path),
        "duration": duration,
    }


def _resolve_speaker(speaker: str, speaker_names: dict = None) -> str:
    """Map a raw speaker label to its display name."""
    if speaker_names and speaker in speaker_names:
        return speaker_names[speaker]
    return speaker


def format_transcript_text(
    result: dict,
    include_timestamps: bool = True,
    speaker_names: dict = None,
) -> str:
    """Format transcription result as readable text."""
    lines = []
    current_speaker = None

    for seg in result["segments"]:
        speaker = _resolve_speaker(seg.get("speaker", "UNKNOWN"), speaker_names)
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        text = seg.get("text", "").strip()

        if not text:
            continue

        if speaker != current_speaker:
            current_speaker = speaker
            lines.append("")  # blank line between speakers
            if include_timestamps:
                lines.append(
                    f"[{format_timestamp(start)} - {format_timestamp(end)}] {speaker}:"
                )
            else:
                lines.append(f"{speaker}:")
            lines.append(f"  {text}")
        else:
            # Same speaker continues — append to their block
            lines.append(f"  {text}")

    return "\n".join(lines).strip()


def format_transcript_srt(result: dict, speaker_names: dict = None) -> str:
    """Format transcription result as SRT subtitles with speaker labels."""
    srt_lines = []
    for i, seg in enumerate(result["segments"], 1):
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        speaker = _resolve_speaker(seg.get("speaker", "UNKNOWN"), speaker_names)
        text = seg.get("text", "").strip()
        if not text:
            continue

        start_ts = format_srt_timestamp(start)
        end_ts = format_srt_timestamp(end)
        srt_lines.append(f"{i}")
        srt_lines.append(f"{start_ts} --> {end_ts}")
        srt_lines.append(f"[{speaker}] {text}")
        srt_lines.append("")

    return "\n".join(srt_lines)


def format_srt_timestamp(seconds: float) -> str:
    """Format seconds as SRT timestamp (HH:MM:SS,mmm)."""
    td = timedelta(seconds=seconds)
    total_seconds = td.total_seconds()
    hours = int(total_seconds // 3600)
    minutes = int((total_seconds % 3600) // 60)
    secs = int(total_seconds % 60)
    millis = int((total_seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def format_transcript_markdown(result: dict, speaker_names: dict = None) -> str:
    """Format transcription result as Markdown."""
    lines = []
    audio_name = Path(result.get("audio_path", "audio")).stem
    duration = format_timestamp(result.get("duration", 0))

    # Header
    lines.append(f"# Transcript: {audio_name}")
    lines.append("")
    lines.append(f"**Duration:** {duration}")

    # Speaker legend
    speakers = sorted(set(
        seg.get("speaker", "UNKNOWN") for seg in result["segments"]
    ))
    if speakers:
        names = [_resolve_speaker(s, speaker_names) for s in speakers]
        lines.append(f"**Speakers:** {', '.join(names)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Body
    current_speaker = None
    for seg in result["segments"]:
        speaker = _resolve_speaker(seg.get("speaker", "UNKNOWN"), speaker_names)
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if not text:
            continue

        if speaker != current_speaker:
            current_speaker = speaker
            lines.append("")
            lines.append(f"**{speaker}** `{format_timestamp(start)}`")
            lines.append("")
            lines.append(text)
        else:
            lines.append(text)

    return "\n".join(lines).strip()


def format_transcript_rtf(result: dict, speaker_names: dict = None) -> str:
    """Format transcription result as RTF."""

    def _rtf_escape(text: str) -> str:
        out = []
        for ch in text:
            code = ord(ch)
            if ch == "\\":
                out.append("\\\\")
            elif ch == "{":
                out.append("\\{")
            elif ch == "}":
                out.append("\\}")
            elif code > 127:
                out.append(f"\\u{code}?")
            else:
                out.append(ch)
        return "".join(out)

    parts = []
    parts.append(r"{\rtf1\ansi\deff0")
    parts.append(r"{\fonttbl{\f0\fswiss Helvetica;}{\f1\fmodern Courier New;}}")
    parts.append(r"{\colortbl;\red0\green0\blue0;\red80\green80\blue80;}")
    parts.append(r"\f0\fs24")
    parts.append("")

    # Title
    audio_name = Path(result.get("audio_path", "audio")).stem
    duration = format_timestamp(result.get("duration", 0))
    parts.append(rf"{{\b\fs32 Transcript: {_rtf_escape(audio_name)}\b0\fs24\par}}")
    parts.append(rf"Duration: {duration}\par")
    parts.append(r"\line")

    current_speaker = None
    for seg in result["segments"]:
        speaker = _resolve_speaker(seg.get("speaker", "UNKNOWN"), speaker_names)
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if not text:
            continue

        if speaker != current_speaker:
            current_speaker = speaker
            parts.append(r"\par")
            parts.append(
                rf"{{\b {_rtf_escape(speaker)}\b0}} "
                rf"{{\f1\fs20\cf2 [{format_timestamp(start)}]\f0\fs24\cf1}}\par"
            )
            parts.append(rf"{_rtf_escape(text)}\par")
        else:
            parts.append(rf"{_rtf_escape(text)}\par")

    parts.append("}")
    return "\n".join(parts)


def save_transcript_docx(result: dict, output_path: str, speaker_names: dict = None):
    """Save transcription result as a Word .docx file."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    audio_name = Path(result.get("audio_path", "audio")).stem
    title = doc.add_heading(f"Transcript: {audio_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    duration = format_timestamp(result.get("duration", 0))
    speakers = sorted(set(
        seg.get("speaker", "UNKNOWN") for seg in result["segments"]
    ))
    names = [_resolve_speaker(s, speaker_names) for s in speakers]
    meta = doc.add_paragraph()
    meta.add_run(f"Duration: {duration}").bold = True
    meta.add_run(f"  |  Speakers: {', '.join(names)}")
    doc.add_paragraph("").paragraph_format.space_after = Pt(4)

    # Body
    current_speaker = None
    for seg in result["segments"]:
        speaker = _resolve_speaker(seg.get("speaker", "UNKNOWN"), speaker_names)
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if not text:
            continue

        if speaker != current_speaker:
            current_speaker = speaker
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            run_name = p.add_run(f"{speaker}  ")
            run_name.bold = True
            run_name.font.size = Pt(11)
            run_ts = p.add_run(f"[{format_timestamp(start)}]")
            run_ts.font.size = Pt(9)
            run_ts.font.color.rgb = RGBColor(120, 120, 120)

            tp = doc.add_paragraph(text)
            tp.paragraph_format.space_before = Pt(2)
            tp.paragraph_format.space_after = Pt(2)
        else:
            tp = doc.add_paragraph(text)
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after = Pt(2)

    doc.save(output_path)
    return output_path


def save_outputs(result: dict, output_dir: str, base_name: str, speaker_names: dict = None):
    """Save transcript in all formats. Returns dict of format -> path."""
    os.makedirs(output_dir, exist_ok=True)
    paths = {}

    # Plain text with timestamps
    txt_path = os.path.join(output_dir, f"{base_name}_transcript.txt")
    with open(txt_path, "w") as f:
        f.write(f"TRANSCRIPT: {result['audio_path']}\n")
        f.write(f"Duration: {format_timestamp(result['duration'])}\n")
        f.write("=" * 60 + "\n\n")
        f.write(format_transcript_text(result, speaker_names=speaker_names))
    paths["txt"] = txt_path

    # JSON (full data with timestamps and word-level detail)
    json_path = os.path.join(output_dir, f"{base_name}_transcript.json")
    with open(json_path, "w") as f:
        json.dump(result, f, indent=2, default=str)
    paths["json"] = json_path

    # SRT subtitles
    srt_path = os.path.join(output_dir, f"{base_name}_transcript.srt")
    with open(srt_path, "w") as f:
        f.write(format_transcript_srt(result, speaker_names=speaker_names))
    paths["srt"] = srt_path

    # Markdown
    md_path = os.path.join(output_dir, f"{base_name}_transcript.md")
    with open(md_path, "w") as f:
        f.write(format_transcript_markdown(result, speaker_names=speaker_names))
    paths["md"] = md_path

    # Rich Text
    rtf_path = os.path.join(output_dir, f"{base_name}_transcript.rtf")
    with open(rtf_path, "w") as f:
        f.write(format_transcript_rtf(result, speaker_names=speaker_names))
    paths["rtf"] = rtf_path

    # Word document
    docx_path = os.path.join(output_dir, f"{base_name}_transcript.docx")
    save_transcript_docx(result, docx_path, speaker_names=speaker_names)
    paths["docx"] = docx_path

    for fmt, p in paths.items():
        print(f"  Saved: {p}")

    return paths


def main():
    parser = argparse.ArgumentParser(
        description="Transcribe audio with speaker diarization (fully local)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic transcription of a hearing recording
  python transcribe.py hearing.m4a --min-speakers 3 --max-speakers 6

  # Use a smaller/faster model
  python transcribe.py hearing.m4a --model small

  # Specify HF token directly
  python transcribe.py hearing.m4a --hf-token hf_xxxxx

  # Output to a specific directory
  python transcribe.py hearing.m4a -o ./transcripts
        """,
    )
    parser.add_argument("audio", help="Path to audio file (m4a, mp3, wav, etc.)")
    parser.add_argument(
        "--model",
        default="large-v3",
        choices=["tiny", "base", "small", "medium", "large-v3"],
        help="Whisper model size (default: large-v3 for best accuracy)",
    )
    parser.add_argument("--device", default=None, help="Device: cuda, mps, cpu (auto-detected)")
    parser.add_argument("--compute-type", default=None, help="Compute type (auto-selected)")
    parser.add_argument("--hf-token", default=None, help="HuggingFace token (or set HF_TOKEN env var)")
    parser.add_argument("--min-speakers", type=int, default=None, help="Min expected speakers")
    parser.add_argument("--max-speakers", type=int, default=None, help="Max expected speakers")
    parser.add_argument("--language", default="en", help="Language code (default: en)")
    parser.add_argument("--batch-size", type=int, default=16, help="Batch size (lower if OOM)")
    parser.add_argument("-o", "--output-dir", default=None, help="Output directory (default: same as audio)")

    args = parser.parse_args()

    audio_path = Path(args.audio)
    if not audio_path.exists():
        print(f"Error: File not found: {audio_path}")
        sys.exit(1)

    output_dir = args.output_dir or str(audio_path.parent)
    base_name = audio_path.stem

    print("=" * 60)
    print("  HEARING TRANSCRIBER")
    print("  Local transcription + speaker diarization")
    print("=" * 60)

    result = transcribe_audio(
        audio_path=str(audio_path),
        model_size=args.model,
        device=args.device,
        compute_type=args.compute_type,
        hf_token=args.hf_token,
        min_speakers=args.min_speakers,
        max_speakers=args.max_speakers,
        language=args.language,
        batch_size=args.batch_size,
    )

    print("\nSaving outputs...")
    paths = save_outputs(result, output_dir, base_name)
    print(f"\nDone! {len(paths)} files saved to {output_dir}")


if __name__ == "__main__":
    main()
